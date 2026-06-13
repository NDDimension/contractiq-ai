import os
from dotenv import load_dotenv
from langchain_community.document_loaders import PyPDFLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
from langchain_groq import ChatGroq
from langchain_core.documents import Document

load_dotenv()

_embeddings = None  # cached so model loads once


def _get_embeddings():
    global _embeddings
    if _embeddings is None:
        _embeddings = HuggingFaceEmbeddings(model_name="all-MiniLM-L6-v2")
    return _embeddings


def _splitter():
    return RecursiveCharacterTextSplitter(
        chunk_size=1000,
        chunk_overlap=150,
        separators=["\n\n", "\n", " ", ""]
    )


def load_contract_pdf(pdf_path: str):
    loader = PyPDFLoader(pdf_path)
    documents = loader.load()
    return _splitter().split_documents(documents)


def load_contract_txt(txt_path: str):
    with open(txt_path, "rb") as f:
        raw = f.read()

    try:
        text = raw.decode("utf-8")
    except UnicodeDecodeError:
        text = raw.decode("latin-1")

    return load_contract_text(text)


def load_contract_docx(docx_path: str):
    try:
        from docx import Document as DocxDocument
    except ImportError as exc:
        raise RuntimeError("DOCX support requires python-docx. Run: pip install -r requirements.txt") from exc

    docx = DocxDocument(docx_path)
    parts = [para.text for para in docx.paragraphs if para.text.strip()]

    for table in docx.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return load_contract_text("\n\n".join(parts))


def load_contract_file(file_path: str, extension: str):
    ext = extension.lower()
    if ext == ".pdf":
        return load_contract_pdf(file_path)
    if ext == ".txt":
        return load_contract_txt(file_path)
    if ext == ".docx":
        return load_contract_docx(file_path)
    raise ValueError("Unsupported file type. Please upload a PDF, TXT, or DOCX file.")


def load_contract_text(text: str):
    doc = Document(page_content=text)
    return _splitter().split_documents([doc])


def build_vectorstore(chunks):
    embeddings = _get_embeddings()
    return FAISS.from_documents(chunks, embeddings)


def get_llm():
    return ChatGroq(
        model="llama-3.3-70b-versatile",
        temperature=0,
        groq_api_key=os.getenv("GROQ_API_KEY")
    )
